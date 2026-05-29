"""
Business Plan Parser — m1_bp_parser.py
Uses Claude to extract investor-matching fields from a BP PDF/DOCX.

Public API (called by demo/app.py — interface unchanged):
    extract_text(file_path: str) -> str
    parse_with_claude(text: str) -> dict

Backwards-compat alias so app.py needs zero changes:
    parse_with_gemini = parse_with_claude

CLI:
    python scraper/m1_bp_parser.py --file path/to/bp.pdf
    python scraper/m1_bp_parser.py --file path/to/bp.docx --output out.json
    python scraper/m1_bp_parser.py --file path/to/bp.pdf --flat
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

import anthropic
from dotenv import load_dotenv

load_dotenv()

# ── Claude client ──────────────────────────────────────────────────────────

_api_key = os.getenv("ANTHROPIC_API_KEY")
if not _api_key:
    sys.exit("ERROR: ANTHROPIC_API_KEY not set in .env")

client = anthropic.Anthropic(api_key=_api_key)
MODEL  = "claude-haiku-4-5-20251001"

# ~30k chars keeps us well within Claude's context while covering most BPs
MAX_CHARS = 30_000


# ── Text extraction (interface unchanged) ──────────────────────────────────

def extract_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages  = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def extract_docx(path: str) -> str:
    from docx import Document
    doc   = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cell_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if cell_text:
                parts.append(cell_text)
    return "\n".join(parts)


def extract_text(file_path: str) -> str:
    """Extract raw text from PDF or DOCX. Raises ValueError on unsupported types."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        text = extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        text = extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type '{ext}'. Use .pdf or .docx")

    text = text.strip()
    if not text:
        raise ValueError("No text extracted — file may be scanned/image-only")

    if len(text) > MAX_CHARS:
        print(f"  ⚠  Truncated {len(text):,} → {MAX_CHARS:,} chars")
        text = text[:MAX_CHARS]

    return text


# ── Prompts ────────────────────────────────────────────────────────────────
#
# Output schema mirrors the investor DB (active.jsonl) so every extracted
# field maps 1-to-1 onto an investor attribute for matching.
#
# Three tiers follow the matching architecture from the design doc:
#   Tier 1  hard-filter fields  (gate: any miss → investor excluded)
#   Tier 2  soft-score fields   (weighted contribution to final rank score)
#   Tier 3  semantic field      (narrative_text → embedding → cosine match)
#
# Confidence tags let the scorer apply penalties:
#   "high"    explicitly stated in the document
#   "medium"  clearly implied / inferable
#   "low"     guessed from weak signals
#   "missing" absent from doc; value will be null

SYSTEM_PROMPT = """\
You are an expert investment analyst extracting structured data from a \
business plan for investor matching.

Output ONLY a single valid JSON object — no markdown fences, no prose, \
no keys outside the schema. Use null for fields that are genuinely absent."""

USER_PROMPT = """\
Extract investor-matching fields from the business plan below.
Every key in the schema must appear in your output.

━━━ SCHEMA ━━━

{{
  "company_name": "string",
  "tagline":      "string — one sentence: what the company does and for whom",

  "tier1": {{

    "funding_stage": {{
      "value":      "one of: pre_seed | seed | series_a | series_b | series_c | growth",
      "raw":        "exact phrase from doc, e.g. 'Raising Series A'",
      "confidence": "high | medium | low | missing"
    }},

    "raise_amount_usd": {{
      "value":      null,
      "value_min":  null,
      "value_max":  null,
      "note":       "raw string from doc e.g. '$3M' or '$2M-$5M round'",
      "confidence": "high | medium | low | missing"
    }},

    "sector_tags": {{
      "value": ["choose from: fintech, ai_ml, ai_infra, healthcare, healthtech,
                 climate, cleantech, deeptech, edtech, saas, consumer,
                 marketplace, hardware, defense, biotech, proptech, legaltech, other"],
      "raw":        "the description phrases you used to classify",
      "confidence": "high | medium | low | missing"
    }},

    "geo": {{
      "hq_city":        "string or null",
      "hq_state":       "string or null",
      "hq_country":     "string or null",
      "remote_first":   false,
      "target_markets": ["geographic markets the product targets"],
      "confidence":     "high | medium | low | missing"
    }},

    "investor_type_preference": {{
      "value":      ["subset of: vc, pe, angel, cvc, accelerator, family_office"],
      "confidence": "high | medium | low | missing"
    }},

    "seeking_lead": {{
      "value":      true,
      "note":       "true if doc says seeking lead / lead investor, or clearly implies it",
      "confidence": "high | medium | low | missing"
    }},

    "instrument": {{
      "value":      "one of: equity | safe | convertible_note | debt | unknown",
      "confidence": "high | medium | low | missing"
    }}
  }},

  "tier2": {{

    "business_model": {{
      "value": ["subset of: saas_subscription, marketplace, deeptech_ip,
                 hardware_software, transaction_fee, usage_based,
                 consumer_app, b2b_services, other"],
      "raw":        "how the company makes money, as stated in doc",
      "confidence": "high | medium | low | missing"
    }},

    "customer_type": {{
      "value":      "b2b | b2c | b2b2c | government | mixed",
      "confidence": "high | medium | low | missing"
    }},

    "traction": {{
      "arr_usd":          null,
      "mrr_usd":          null,
      "gmv_usd":          null,
      "paying_customers": null,
      "pilots":           null,
      "users":            null,
      "pre_revenue":      false,
      "note":             "verbatim traction statement from doc",
      "confidence":       "high | medium | low | missing"
    }},

    "team_background": {{
      "value": ["subset of: ex_faang, repeat_founder, phd_technical,
                 domain_expert, ex_finance, ex_consulting, first_time_founder"],
      "founders": [
        {{"name": "string or null", "title": "string or null", "note": "string or null"}}
      ],
      "confidence": "high | medium | low | missing"
    }},

    "use_of_funds": {{
      "value":      "how raised capital will be deployed (hiring, product, GTM…)",
      "confidence": "high | medium | low | missing"
    }},

    "sub_sector_tags": {{
      "value": ["从 BP 核心技术和产品描述中提取 3–6 个最能代表
                 该公司细分方向的标签。只能从以下列表选择，
                 不得使用列表之外的词：

        医疗健康类：
        digital_health, medtech, drug_discovery, platform_biotech,
        clinical_ai, diagnostics, immunotherapy, oncology, genomics,
        synthetic_biology, cell_therapy, gene_therapy, ai_healthcare,
        ai_drug_discovery, mental_health, neuroscience, rare_disease,
        regenerative_medicine, cardiovascular, longevity, bioelectronics,
        musculoskeletal, ophthalmology, womens_health

        气候清洁能源类：
        energy_storage, carbon_capture, agtech_climate, sustainable_materials,
        grid_infra, circular_economy, ev_charging, green_hydrogen,
        sustainable_food, solar, water_tech, nuclear, building_efficiency,
        wind, carbon_markets, advanced_manufacturing

        AI和机器人类：
        robotics_ai, ai_agents, foundation_models, generative_ai,
        ai_chips_infra, autonomous_vehicles, industrial_robotics,
        computer_vision, ai_for_science, ai_security, mlops,
        ai_finance, ai_legal, edge_ai, synthetic_data, nlp

        深科技硬件类：
        space_tech, semiconductors, advanced_materials, quantum_computing,
        sensors, photonics, drones, satellite_comms, nanotechnology

        国防安全类：
        autonomous_defense, cybersecurity, cybersecurity_saas,
        military_ai, dual_use_tech, counter_drone

        金融科技类：
        payments, b2b_payments, lending, crypto_web3, insurtech,
        embedded_finance, wealthtech, neobank, financial_infra,
        real_estate_finance, regtech, accounting_tax, defi

        企业软件类：
        devtools, vertical_saas, data_infra, supply_chain, hr_tech,
        legaltech, legal_tech, api_first, marketing_tech, cloud_infra

        消费类：
        d2c, edtech, consumer_health, marketplace_consumer,
        creator_economy, proptech, mobility, precision_agriculture,
        agtech, pet_tech

        其他：
        generalist"],
      "confidence": "high | medium | low | missing"
    }}
  }},

  "tier3": {{
    "narrative_text": "关键叙事段落，涵盖 problem/solution/why now，最多 300 词，文本内的双引号替换为单引号。"
  }},

  "missing_signals": ["important investor-matching fields absent from the doc"],

  "overall_confidence": 0.0
}}

━━━ RULES ━━━
- overall_confidence: float 0.0-1.0. 1.0 = all Tier 1 + Tier 2 fields present and clear.
- raise_amount_usd: parse to integer USD (3000000, not "3M"). Use value_min/value_max for ranges.
- sector_tags.value: multiple values allowed. Map jargon to the closest enum term.
- traction.pre_revenue: set true only when doc explicitly says pre-revenue or has zero revenue signals.
- narrative_text: must be original text from the document, not a paraphrase or summary.
- Do not add any key not in the schema above.

━━━ BUSINESS PLAN TEXT ━━━
{text}
━━━ END ━━━"""


# ── Claude API call ────────────────────────────────────────────────────────

def parse_with_claude(text: str) -> dict:
    prompt = USER_PROMPT.format(text=text)

    for attempt in range(3):
        try:
            response = client.messages.create(
                model=MODEL,
                max_tokens=2500,
                temperature=0.1,
                system=SYSTEM_PROMPT,
                messages=[{"role": "user", "content": prompt}],
            )

            raw = response.content[0].text.strip()

            print("=== RAW CLAUDE OUTPUT ===")
            print(raw)
            print("=== END RAW OUTPUT ===")

            # Strip markdown fences
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```\s*$", "", raw)

            # 方法1：直接解析
            try:
                return json.loads(raw)
            except json.JSONDecodeError:
                pass

            # 方法2：提取最外层 {} 之间的内容
            match = re.search(r'\{.*\}', raw, re.DOTALL)
            if match:
                try:
                    return json.loads(match.group(0))
                except json.JSONDecodeError:
                    pass

            # 方法3：用 json_repair 修复（如果安装了的话）
            try:
                import json_repair
                return json_repair.loads(raw)
            except (ImportError, Exception):
                pass

            # 全部失败，打印原始内容调试
            print(f"  Raw response (attempt {attempt+1}/3):")
            print(raw[:1000])
            print("  ...")
            raise json.JSONDecodeError("All parse methods failed", raw, 0)

        except json.JSONDecodeError as e:
            print(f"  JSON parse error (attempt {attempt+1}/3): {e}")
            if attempt == 2:
                raise RuntimeError(f"Claude returned invalid JSON after 3 attempts: {e}") from e

        except anthropic.APIStatusError as e:
            print(f"  Claude API error (attempt {attempt+1}/3): HTTP {e.status_code}")
            if e.status_code in (401, 403):
                raise RuntimeError("Invalid ANTHROPIC_API_KEY") from e
            if e.status_code == 429:
                raise RuntimeError("Claude rate limit hit") from e
            if attempt == 2:
                raise

        except Exception as e:
            print(f"  Unexpected error (attempt {attempt+1}/3): {str(e)[:120]}")
            if attempt == 2:
                raise

    raise RuntimeError("All Claude attempts failed")


# Backwards-compat alias — app.py calls parse_with_gemini(); no edits needed there
parse_with_gemini = parse_with_claude


# ── Flat display helper ────────────────────────────────────────────────────

def flatten_profile(profile: dict) -> dict:
    """
    Collapse the nested bp_profile into a flat dict for the M1 frontend.
    The full nested profile is preserved under '_profile' for the matching engine.
    """
    t1       = profile.get("tier1", {})
    t2       = profile.get("tier2", {})
    t3       = profile.get("tier3", {})
    traction = t2.get("traction", {})

    def v(d):
        """Extract .value from a confidence-tagged field."""
        return d.get("value") if isinstance(d, dict) else d

    def join_list(d):
        """Turn a list .value into a comma-separated string."""
        val = v(d)
        if isinstance(val, list):
            return ", ".join(str(x) for x in val) if val else None
        return val

    def fmt_usd(n):
        if n is None:
            return None
        return f"${n / 1_000_000:.1f}M" if n >= 1_000_000 else f"${n:,}"

    # Format raise amount
    amt    = t1.get("raise_amount_usd", {}) or {}
    lo     = amt.get("value_min")
    hi     = amt.get("value_max")
    single = amt.get("value")
    if lo and hi:
        capital_need = f"{fmt_usd(lo)}–{fmt_usd(hi)}"
    elif single:
        capital_need = fmt_usd(single)
    else:
        capital_need = amt.get("note")

    # Format HQ location
    geo = t1.get("geo", {}) or {}
    geo_parts = [geo.get("hq_city"), geo.get("hq_state"), geo.get("hq_country")]
    geo_str   = ", ".join(p for p in geo_parts if p) or None

    return {
        # ── Tier 1
        "company_name":             profile.get("company_name"),
        "tagline":                  profile.get("tagline"),
        "funding_stage":            v(t1.get("funding_stage", {})),
        "capital_need":             capital_need,
        "sector":                   join_list(t1.get("sector_tags", {})),
        "geography":                geo_str,
        "investor_type_preference": join_list(t1.get("investor_type_preference", {})),
        "seeking_lead":             v(t1.get("seeking_lead", {})),
        "instrument":               v(t1.get("instrument", {})),
        # ── Tier 2
        "business_model":           join_list(t2.get("business_model", {})),
        "customer_type":            v(t2.get("customer_type", {})),
        "arr_usd":                  traction.get("arr_usd")          if isinstance(traction, dict) else None,
        "mrr_usd":                  traction.get("mrr_usd")          if isinstance(traction, dict) else None,
        "paying_customers":         traction.get("paying_customers") if isinstance(traction, dict) else None,
        "pre_revenue":              traction.get("pre_revenue", False) if isinstance(traction, dict) else None,
        "key_traction":             traction.get("note")             if isinstance(traction, dict) else None,
        "team_background":          join_list(t2.get("team_background", {})),
        "use_of_funds":             v(t2.get("use_of_funds", {})),
        "sub_sector_tags":          join_list(t2.get("sub_sector_tags", {})),
        # ── Tier 3
        "narrative_text":           t3.get("narrative_text"),
        # ── Meta
        "overall_confidence":       profile.get("overall_confidence", 0.0),
        "missing_signals":          profile.get("missing_signals", []),
        # Full nested profile for the matching engine
        "_profile":                 profile,
    }


# ── CLI entry point ────────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser(
        description="Parse a BP PDF/DOCX with Claude and output structured JSON"
    )
    ap.add_argument("--file",   required=True, help="Path to .pdf or .docx")
    ap.add_argument("--output", default="output/business_profile.json",
                    help="Output path (default: output/business_profile.json)")
    ap.add_argument("--flat",   action="store_true",
                    help="Also print a flattened view (without the nested blob)")
    args = ap.parse_args()

    if not os.path.exists(args.file):
        sys.exit(f"ERROR: File not found: {args.file}")

    print(f"📄 Reading:          {args.file}")
    text = extract_text(args.file)
    print(f"   Extracted {len(text):,} chars")

    print(f"🤖 Calling Claude ({MODEL})…")
    profile = parse_with_claude(text)

    # Attach run metadata
    profile["_source_file"] = os.path.basename(args.file)
    profile["_parsed_at"]   = datetime.now(timezone.utc).isoformat()

    # Save full nested output
    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(profile, f, ensure_ascii=False, indent=2)
    print(f"\n✅  Saved → {out_path}\n")

    if args.flat:
        flat = flatten_profile(profile)
        flat.pop("_profile", None)
        print("── Flattened view ──")
        print(json.dumps(flat, ensure_ascii=False, indent=2))
    else:
        print(json.dumps(profile, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
